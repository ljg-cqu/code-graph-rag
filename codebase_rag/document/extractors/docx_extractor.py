"""DOCX document extractor.

Optional extractor that requires python-docx.
"""

from __future__ import annotations

import asyncio
from datetime import UTC, datetime
from pathlib import Path

from .base import (
    BaseDocumentExtractor,
    ExtractedDocument,
    ExtractedSection,
)
from ..error_handling import ErrorType, ExtractionException


class DocxExtractor(BaseDocumentExtractor):
    """
    Extractor for Microsoft Word files (.docx).

    Requires python-docx for text extraction.
    """

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, file_path: Path) -> ExtractedDocument:
        """Extract content from DOCX file."""
        # Validate path
        validated_path = self._validate_path(file_path)

        # Check file size
        max_size_mb = self.get_config("max_file_size_mb", 50)
        file_size_mb = validated_path.stat().st_size / (1024 * 1024)
        if file_size_mb > max_size_mb:
            raise ExtractionException(
                path=str(file_path),
                error_type=ErrorType.FILE_TOO_LARGE,
                message=f"File size ({file_size_mb:.1f}MB) exceeds limit ({max_size_mb}MB)",
            )

        # Try to import DOCX library
        try:
            from docx import Document  # type: ignore
        except ImportError:
            raise ExtractionException(
                path=str(file_path),
                error_type=ErrorType.MISSING_DEPENDENCY,
                message="DOCX extraction requires python-docx. Install with: uv add python-docx",
            )

        return self._extract_with_docx(validated_path, file_path, Document)

    def _extract_with_docx(
        self, validated_path: Path, original_path: Path, Document: type
    ) -> ExtractedDocument:
        """Extract using python-docx."""
        doc = Document(validated_path)

        content_parts: list[str] = []
        sections: list[ExtractedSection] = []
        code_blocks: list[str] = []

        # Extract paragraphs and track headings
        current_section: ExtractedSection | None = None
        section_content: list[str] = []
        line_num = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading
            style_name = para.style.name.lower() if para.style else ""
            is_heading = "heading" in style_name or "title" in style_name

            if is_heading:
                # Save previous section
                if current_section and section_content:
                    current_section.content = "\n".join(section_content)
                    current_section.end_line = line_num - 1
                    sections.append(current_section)

                # Start new section
                level = 1
                if "heading 1" in style_name:
                    level = 1
                elif "heading 2" in style_name:
                    level = 2
                elif "heading 3" in style_name:
                    level = 3
                elif "heading" in style_name:
                    # Extract number from "Heading X"
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1

                current_section = ExtractedSection(
                    title=text,
                    level=level,
                    start_line=line_num,
                    end_line=line_num,
                    content="",
                    subsections=[],
                )
                section_content = []
            else:
                content_parts.append(text)
                section_content.append(text)

            line_num += 1

        # Save last section
        if current_section and section_content:
            current_section.content = "\n".join(section_content)
            current_section.end_line = line_num - 1
            sections.append(current_section)

        # Extract code blocks from tables (sometimes code is in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and len(cell_text) > 20:
                        # Heuristic: if it looks like code, add it
                        if any(
                            kw in cell_text
                            for kw in ["def ", "class ", "function ", "import ", "const "]
                        ):
                            code_blocks.append(cell_text)

        content = "\n\n".join(content_parts)
        modified_date = datetime.fromtimestamp(
            validated_path.stat().st_mtime, UTC
        ).isoformat()

        return ExtractedDocument(
            path=str(original_path),
            file_type=".docx",
            content=content,
            sections=sections,
            code_blocks=code_blocks,
            code_references=self._extract_code_references(content),
            word_count=len(content.split()),
            modified_date=modified_date,
        )

    async def extract_async(self, file_path: Path) -> ExtractedDocument:
        """Async extraction for DOCX files."""
        return await asyncio.to_thread(self.extract, file_path)


__all__ = ["DocxExtractor"]